__author__ = 'swallow'
__language__= 'python 3.0'

from functools import total_ordering
import os
from pickle import FALSE
import sys
import time
#from datetime import timedelta
import datetime

import locale

import json
import random
from tkinter.filedialog import SaveAs
import docx
import requests
#from bs4 import BeautifulSoup
import re

from selenium import webdriver

import pythoncom
import win32com
from win32com.client import Dispatch,DispatchEx
from PIL import ImageGrab, Image
import pytesseract
import keyboard

import fitz

import CommonFunc
from CommonFunc import isEmptyValue
from CommonFunc import copyspecFile
from CommonFunc import getCellValueinString
from CommonFunc import rgb_to_hex
from CommonFunc import findFile

from CommonFunc import nodeTree
from CommonFunc import nodeObj


from redminelib import managers
from redminelib import Redmine 
#from polarion import polarion

from docx import Document

from PyPDF2 import PdfFileReader

from pdf2docx import parse

from typing import Tuple

from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer

import filecmp

from atlassian import Confluence

#import pymysql as MySqldb

#import myDB 
try:
    import xml.etree.cElementTree as ET
except ImportError:
    import xml.etree.ElementTree as ET

import ocrmypdf
from jira import JIRA

header = {

    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',

    'Accept-Encoding': 'gzip, deflate, sdch',

    'Accept-Language': 'zh-CN,zh;q=0.8',

    'Connection': 'keep-alive',

    'User-Agent': 'Mozilla/5.0 (Windows NT 6.3; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/43.0.235'

}

timeout = random.choice(range(80, 180))



class  ToolsFixer(object):
    result_list = []
    mulity_state_list = []
    shapes_list0 = {}
    shapes_list1 = {}

    def __init__(self,path='.'):    
        self._path=path
        self.abspath=os.path.abspath(self._path) # 默认当前目录
        self.excelApp = DispatchEx("Excel.Application")
        self.wordApp = DispatchEx('Word.Application')
        self.pptApp = DispatchEx('PowerPoint.Application')

        self.wordApp.visible = True
        self.excelApp.visible = True
        self.pptApp.visible = True

        self.excelApp.DisplayAlerts = False

        self.excelApp.Workbooks.Add()
        self.resultBook = self.excelApp.ActiveWorkBook


    def __exit_(self, *args):
        return
        #self.resultBook.Close(SaveChanges = 1)
        self.excelApp.Application.Quit()
        self.wordApp.Application.Quit()
        self.pptApp.Application.Quit()


    def testJira(self):                                                                                                                                          
        # Specify a server key. It is your  domain 
        # name link.
        #For Lomami
        '''
        jiraOptions = {'server': "https://jira.vplatf.com"}
        # Get a JIRA client instance, Pass 
        # Authentication parameters
        # and  Server name.
        # emailID = your emailID
        # token = token you receive after registration
        jira = JIRA(options = jiraOptions, 
                    basic_auth = ("iAuto001","Nd5/WC#h"))
        '''
        jiraOptions = {'server': "https://tmecjira.pjt.nedl.com.cn", 'verify':False}
        # Get a JIRA client instance, Pass 
        # Authentication parameters
        # and  Server name.
        # emailID = your emailID
        # token = token you receive after registration
        jira = JIRA(options = jiraOptions,
                    basic_auth = ("huangxiaoyan","24MMSakura"))


        # While fetching details of a single issue,
        # pass its UniqueID or Key.
        
        '''
        #Get project
        project = jira.projects("QASIDS")
        
        #Create issue
        issue_dict = {
            'project': {'id': 15303},
            'summary': '[USB] 15.6.5.3 ランダム再生 履歴の確認',
            'description': '10-1.制御仕様書(オーディオ・メディア)_220428.docx 15.6.5.3 ランダム再生　により、履歴は指定範囲内の全曲を対象とする。',
            'issuetype': {'id': '10300'},
            'customfield_19620':{'value':'iAUTO'},
            'customfield_19621':{'value':'AISIN'},
            'reporter':{'key':'JIRAUSER24317','name':'iAuto001'}
        }
        new_issue = jira.create_issue(fields=issue_dict)
        print('{}: {}:{}'.format(new_issue.key,
                                 new_issue.fields.summary,
                                 new_issue.fields.reporter.displayName))    
        #Add comment
        #comment = jira.add_comment('QASIDSIA-169', '社内確認中')
        '''

        self.resultBook.ActiveSheet.Name = "JIRA status"
        resultSheet = self.resultBook.ActiveSheet
        resultSheet.Cells(1,1).Value = "JIRA NO."
        resultSheet.Cells(1,2).Value = "Summary"
        resultSheet.Cells(1,3).Value = "Reporter"
        resultSheet.Cells(1,4).Value = "Assignee"
        resultSheet.Cells(1,5).Value = "Comment body"
        resultSheet.Cells(1,6).Value = "Comment author"
        resultSheet.Cells(1,7).Value = "Comment date"

        wRow = 2
        #Searching issues
        for proj in ("QASIDSHSIA","QASIDSIA"):
            issues_in_proj = jira.search_issues('project=' + proj) #QASIDSHSIA   QASIDSIA
            total_number = issues_in_proj.total
            for i in range(0, int(total_number/50) +1):
                issues_in_proj = jira.search_issues('project=' + proj, startAt=50*i)
                for singleIssue in issues_in_proj:

                    comments_list = jira.comment(singleIssue.key,"").comments    
                    if len(comments_list) > 0:
                        comment = comments_list[len(comments_list)-1]
                        comment_date = comment.updated[0:10].replace("-","")
                        current_date = datetime.datetime.now().strftime('%Y%m%d')
                        #if int(comment_date) >= int(current_date) -1:
                        resultSheet.Cells(wRow,1).Value = singleIssue.key
                        resultSheet.Cells(wRow,2).Value = singleIssue.fields.summary
                        resultSheet.Cells(wRow,3).Value = singleIssue.fields.reporter.displayName
                        
                        if  singleIssue.fields.assignee is not None:
                            resultSheet.Cells(wRow,4).Value = singleIssue.fields.assignee.displayName
                        resultSheet.Cells(wRow,5).Value = comment.body
                        resultSheet.Cells(wRow,6).Value = comment.updateAuthor.displayName
                        resultSheet.Cells(wRow,7).Value = comment.updated
                        resultSheet.Cells(wRow,8).Value = singleIssue.fields.status.name
                        resultSheet.Cells(wRow,9).Value = comment.id
                        wRow += 1
                    else:
                        resultSheet.Cells(wRow,1).Value = singleIssue.key
                        resultSheet.Cells(wRow,2).Value = singleIssue.fields.summary
                        resultSheet.Cells(wRow,3).Value = singleIssue.fields.reporter.displayName
                        if singleIssue.fields.assignee is not None:
                            resultSheet.Cells(wRow,4).Value = singleIssue.fields.assignee.displayName
                        wRow += 1

        self.resultBook.SaveAs(os.getcwd() + "\\Jira_update.xlsx")
        

    def createRedmineTask(self):
        funcfile = os.getcwd() + "\\doc\\redminetask.xlsx"
        funcfileWb  = self.excelApp.Workbooks.Open(funcfile)
        funclistSheet = funcfileWb.Sheets('Sheet2')
        redmine = Redmine('http://iredmine.storm', key='5feb93efacf00e78e74b7a6dbea4d001257544a7', raise_attr_exception = False)

        info = funclistSheet.UsedRange
        nrows = info.Rows.Count

        row = 9
        while row <= nrows:
            #cellValue = getCellValueinString(specSheet.Cells(row, 3))
            if isEmptyValue(funclistSheet.Cells(row, 1)): 
                func_name = getCellValueinString(funclistSheet.Cells(row, 4))

                project_name = 'Tarim_Task'
                project_id = 4086
                subject_str = func_name + ' 请针对航盛的修改做最后的Review'
                parent_id = ''
                assigned_id = getCellValueinString(funclistSheet.Cells(row, 5))
                description = "文件路径：https://svn.ci.iauto.com/svn/smartauto/Tarim_CCS/03_Spec/03_Func/航盛做成/"+func_name + "\n"
                description += "请开发组确认一下指摘的修改是否还有问题，如果没有问题的话，航盛侧要做一个fix版"
                start_date = '2021-12-29'
                due_date = '2021-12-30'

                newissue_id =  self.testRedmine(redmine, project_name, project_id, parent_id, assigned_id, subject_str, description,start_date, due_date)
                funclistSheet.Cells(row, 1).Value = "#" + str(newissue_id)
            row += 1

        funcfileWb.Close(SaveChanges =1)
    
    def createRedmineSubject(self):
        funcfile = os.getcwd() + "\\doc\\redminetask.xlsx"
        funcfileWb  = self.excelApp.Workbooks.Open(funcfile)
        funclistSheet = funcfileWb.Sheets('Sheet2')
        redmine = Redmine('http://iredmine.storm', key='5feb93efacf00e78e74b7a6dbea4d001257544a7', raise_attr_exception = False)

        info = funclistSheet.UsedRange
        nrows = info.Rows.Count

        row = 9
        while row <= nrows:
            #cellValue = getCellValueinString(specSheet.Cells(row, 3))
            if isEmptyValue(funclistSheet.Cells(row, 1)):
                func_name = getCellValueinString(funclistSheet.Cells(row, 7))
                project_name = getCellValueinString(funclistSheet.Cells(row, 2))
                project_id = getCellValueinString(funclistSheet.Cells(row, 3))
                subject_str = '['+ func_name +']' +  getCellValueinString(funclistSheet.Cells(row, 4))
                parent_id = getCellValueinString(funclistSheet.Cells(row, 5))
                assigned_id = getCellValueinString(funclistSheet.Cells(row, 6))
                description = getCellValueinString(funclistSheet.Cells(row, 8))
                start_date = getCellValueinString(funclistSheet.Cells(row, 9))
                due_date = getCellValueinString(funclistSheet.Cells(row, 10))

                newissue_id =  self.testRedmine(redmine, project_name, project_id, parent_id, assigned_id, subject_str, description,start_date, due_date)
                funclistSheet.Cells(row, 1).Value = "#" + str(newissue_id)
            row += 1

        funcfileWb.Close(SaveChanges =1)


    def updateRedmineStatus(self):
        funcfile = os.getcwd() + "\\doc\\RedmineMemberList.xlsx"
        funcfileWb  = self.excelApp.Workbooks.Open(funcfile)
        funclistSheet = funcfileWb.Sheets('1214')
        redmine = Redmine('http://iredmine.storm', key='5feb93efacf00e78e74b7a6dbea4d001257544a7', raise_attr_exception = False)

        info = funclistSheet.UsedRange
        nrows = info.Rows.Count

        row = 2
        while row <= nrows:
            #cellValue = getCellValueinString(specSheet.Cells(row, 3))
            if not isEmptyValue(funclistSheet.Cells(row, )):
                issue_id = getCellValueinString(funclistSheet.Cells(row, 9)).replace("#","")
                issue = redmine.issue.get(issue_id)
                
                funclistSheet.Cells(row, 10).Value = issue.status.name

            row += 1

        funcfileWb.Close(SaveChanges =1)

    def createNormalSpecChangeTickets(self):
            #修改区域=============
            title_str = "【Jira NO 2814】【主】氛围灯音乐联动式样变更"
            parent_id  = '129252'
            start_date = "2023-09-04"
            due_date = "2023-09-11"
            user_name = "hxy"
            #修改区域=============

            SL_list = {
                "hxy": 145,          
                "sunhong": 367,      
                "jiangtian": 427,      
                "yudan": 697
            }
            user_id = SL_list[user_name] 
            member_list = {
                    "_需求组_变更讨论": user_id,      #SL
                    "_FW_SY_变更讨论":156,              #FW_那晓悝
                    "_BCCC HMI_变更讨论" :275,          #HMI 陈猛
                    "_HMI SY_变更讨论" :709,            #HMI 穆坤
                    "_UIFW_变更讨论":387,               #UIFW顾威
                    "_Vehicle_变更讨论": 270,           #Vehicle郭绍峰
                    "_Agent_变更讨论": 248,             #Agent_齐国明
                    "_Telema_变更讨论":183,             #Telema&Update_SY张杰
                    "_Media_变更讨论":333,              #Media尹建成
                    "_Connectivity_变更讨论":339,       #Connectivity徐军军
                    "_Update_变更讨论": 239,            #FW吴庆银
                    "_PVM/RVC_变更讨论":471,            #Update_刘全
                    "_VK_变更讨论":209,                 #控件 张雪宜
                    "_RSE_变更讨论":124,                #Media余兵
                    "_Local Media HMI_SY_变更讨论":709, #穆坤
                    "_FW _变更讨论": 373,               #W王虎成
                    "_Bigdata MediaFW_变更讨论":138,    #李春芳
                    "_GNSS_变更讨论":  330,             #J江汉
                    "_手机连携_变更讨论":  62,          #T_汤文君
                    "_Diag_变更讨论":   376,            #Diag 詹绪文
                    "_Remote parking_变更讨论": 408 ,  #AndroidFW张雪涛
                    "_Meter连携_变更讨论": 471,         # 刘全
                    "_Design_变更讨论": 194,            #Design小伟
                    "_RC_变更讨论":124,                 # Media余兵
                    "_HMI_Launcher_变更讨论":408 ,      #  AndroidFW张雪涛
                    "_Search_变更讨论": 408,            # AndroidFW张雪涛
                    "_Security_变更讨论": 62,           # T_汤文君
                    "_ARACOM_变更讨论": 208 ,           #X夏欢
                    "_合格性测试_变更讨论": 466,        #Z张敏鸣
                    "_系统测试_变更讨论": 466 ,         #Z张敏鸣                    
                    }


            #变更检讨
            description = "※不需要对应时，填写：【机能G名】无Cost\n"
            description = "※和机能组无关时：【机能G名】无关联\n"
            description += "#和机能G无关的情况式样变更对应阶段不再起子票\n"
            description += "但如发现后续更新的式样书变更履历中有对应的变连票，请主动找SL申请追加起对应子票。\n"
            description += "检讨结果填写格式如下：\n"
            description += "【机能G名】\n"
            description += "----对应可否：\n"
            description += "XXXXX\n"
            description += "----对应条件：\n"
            description += "・对应日程\n"
            description += "例如，从式样fix到测试完了需要多久？\n"
            description += "・对应前提\n"
            description += "XXXXX\n"
            description += "----对应内容（予定）：\n"
            description += "１．XXXXX变更点1\n"
            description += "・作业内容：\n"
            description += "XXXXXX\n"
            description += "・设计： XXX人日\n"
            description += "・Coding：XXX人日\n"
            description += "・Test：XXX人日\n"
            description += "・想定作业行数（或文档页数）：XXXX行\n"
            description += "２．XXXXX变更点2\n"
            description += "・作业内容：\n"
            description += "XXXXXX\n"
            description += "・设计： XXX人日\n"
            description += "・Coding：XXX人日\n"
            description += "・Test：XXX人日\n"
            description += "・想定作业行数（或文档页数）：XXXX行\n"
            description += "・对应可能时期：XXXXX日/阶段\n"
            description += "・工数(見積)：XXXXHr （总工数换算成小时）\n"
            description += "----見積作业实际\n"
            description += "确认式样书：XXXHr （检讨花费的工数）\n"

            proj_name = "rhine_change_request_management_toyota"
            proj_id = 1327 #"rhine_change_request_management_toyota","name":"Rhine_变更管理_丰田"       
            track_id = 13 #变更讨论

            self.createiAutoRedmineSubject(proj_name, proj_id, track_id, parent_id, title_str, start_date, due_date, description, member_list, user_name)


    def createiAutoRedmineSubject(self, proj_name, proj_id, track_id, parent_id, title_str, start_date, due_date, description, member_list, user_name):
        user_id = 145
        redmine = Redmine('http://iredmine.iauto.com/', key='d27046e8d2a28cee929411b668c25da4d47453b3', raise_attr_exception = False)
        if user_name == "hxy" :
            user_id = 145
            redmine = Redmine('http://iredmine.iauto.com/', key='d27046e8d2a28cee929411b668c25da4d47453b3', raise_attr_exception = False)
        elif user_name == "sunhong":
            redmine = Redmine('http://iredmine.iauto.com/', key='582be232eabc6a850ef23d69d98117a1a618e792', raise_attr_exception = False) 
            user_id = 367
        elif user_name == "jiangtian":
            redmine = Redmine('http://iredmine.iauto.com/', key='cd8e75a8eed0e5777759d6fbe1993c73f58dada0', raise_attr_exception = False) 
            user_id = 427
        elif user_name == "yudan":
            redmine = Redmine('http://iredmine.iauto.com/', key='b08aeeef66dd0563bad3d4f23d9192b97121b0e5', raise_attr_exception = False) 
            user_id = 697

        for key in member_list:
            c_subject_str = title_str + key
            assigned_id = member_list[key]

            newissue_id =  self.testRedmine(redmine, proj_name, proj_id, parent_id, assigned_id, c_subject_str, description, start_date, due_date, track_id)
            print(newissue_id)
            

    def testRedmine(self, redmine, project_name, project_id, parent_id, assigned_id, subject, description, start_date, due_date,tracker_id):
        # 这里连接redmine，没什么好说的，自己部署的url，账号和密码，照着套就行了  
        projects = redmine.project.all()

        users = redmine.user.all(offset=10, limit=100)

        project = redmine.project.get(project_name)
        new_issue = redmine.issue.new()

        new_issue.project_id=project_id
        new_issue.subject=subject
        new_issue.is_private = 0
        new_issue.tracker_id=tracker_id
        new_issue.description=description
        new_issue.status_id=1
        new_issue.priority_id=2
        new_issue.assigned_to_id=int(assigned_id)
        new_issue.category_id=""
        if parent_id != "":
            new_issue.parent_issue_id=int(parent_id)
        new_issue.start_date = start_date
        new_issue.due_date = due_date
        new_issue.estimated_hours=0
        new_issue.done_ratio=0
        new_issue.fixed_version_id = 159
        new_issue.custom_fields=[{'id': 2, 'value': '1.00'} ]#,{'id': 229, 'value': '需求分析(SL)'},{'id': 66, 'value': 'QASIDSIA-401'},{'id': 122, 'value': '技术Q&A'}]
        #new_issue.custom_fields = [{'id': 10, 'value': '1.00'},{'id': 229, 'value': '需求分析(SL)'}, {'id': 122, 'value': '需求Q&A(需求不明)'}]       

        new_issue.save()

        return new_issue.id


    def updateiAutoRedmineSubject(self):
        # hxy
        user_id = 145
        redmine = Redmine('http://iredmine.iauto.com/', key='d27046e8d2a28cee929411b668c25da4d47453b3', raise_attr_exception = False)

        #jinwan
        #redmine = Redmine('http://iredmine.iauto.com/', key='757f2035a324a88165543baf549c11f60d928936', raise_attr_exception = False) 
        #user_id = 164

        #sunhong
        #redmine = Redmine('http://iredmine.iauto.com/', key='582be232eabc6a850ef23d69d98117a1a618e792', raise_attr_exception = False) 
        #user_id = 367
        issues = redmine.issue.filter(parent_id = 129252)
        member_list = {
                "_需求组_变更讨论": user_id,      #SL
                "_FW_SY_变更讨论":156,              #FW_那晓悝
                "_BCCC HMI_变更讨论" :275,          #HMI 陈猛
                "_HMI SY_变更讨论" :709,            #HMI 穆坤
                "_UIFW_变更讨论":387,               #UIFW顾威
                "_Vehicle_变更讨论": 270,           #Vehicle郭绍峰
                "_Agent_变更讨论": 248,             #Agent_齐国明
                "_Telema_变更讨论":183,      #Telema&Update_SY张杰
                "_Media_变更讨论":333,              #Media尹建成
                "_Connectivity_变更讨论":339,       #Connectivity徐军军
                "_Update_变更讨论": 239,            #FW吴庆银
                "_PVM/RVC_变更讨论":471,            #Update_刘全
                "_VK_变更讨论":209,                 #控件 张雪宜
                "_RSE_变更讨论":124,                #Media余兵
                "_Local Media HMI_SY_变更讨论":709, #穆坤
                "_FW _变更讨论": 373,               #W王虎成
                "_Bigdata MediaFW_变更讨论":138,    #李春芳
                "_GNSS_变更讨论":  330,             #J江汉
                "_手机连携_变更讨论":  62,          #T_汤文君
                "_Diag_变更讨论":   376,            #Diag 詹绪文
                "_Remote parking_变更讨论": 408 ,  #AndroidFW张雪涛
                "_Meter连携_变更讨论": 282,         # H胡乔寿
                "_Design_变更讨论": 117,            #Design孙豆豆
                "_RC_变更讨论":124,                 # Media余兵
                "_HMI_Launcher_变更讨论":408 ,      # AndroidFW张雪涛
                "_Search_变更讨论": 408,            # AndroidFW张雪涛
                "_Security_变更讨论": 62,           # T_汤文君
                "_ARACOM_变更讨论": 208 ,           #X夏欢
                "_合格性测试_变更讨论": 466,        #Z张敏鸣
                "_系统测试_变更讨论": 466 ,         #Z张敏鸣                    
                }
        
        for issue in issues:
            #title_str = issue.subject.replace("追加敏感信息的同意状态也实时上传到中心","追加隐私总条款及敏感信息的状态实时上传到中心")
            for group_name in member_list:
                if issue.subject.find(group_name) > 0 :
                    issue.assigned_to_id = member_list[group_name]
                    issue.due_date = "2023-09-08"
                    #issue.subject = issue.assigned_to_id
                    issue.custom_fields=[{'id': 311, 'value': '外部原因'},{'id': 312, 'value': '有D0通信式样书新input，重新估算变更'} ]
                    issue.save()
                    print(issue.id)
                    break


    def updateTeslinRedmineSubject(self):
        tickets_file = os.getcwd() + "\\doc\\24MM_redmine_infor.xlsx"
        tickets_wb  = self.excelApp.Workbooks.Open(tickets_file)
        tickets_sheet = tickets_wb.Sheets('My tickets') 

        startRow = 2
        info = tickets_sheet.UsedRange
        nrows = info.Rows.Count
        tickets_list = {}
        while startRow <= nrows:
            if not isEmptyValue(tickets_sheet.Cells(startRow,1)):
                tickets_list[getCellValueinString(tickets_sheet.Cells(startRow, 1))] = startRow
            startRow += 1

        wRow = startRow
        redmine = Redmine('http://iredmine.iauto.com/', key='d27046e8d2a28cee929411b668c25da4d47453b3', raise_attr_exception = False)
 
        for proj_id in ("rhine",\
                        "rhine_client_input_management",\
                        "rhine_subject_management", \
                        "rhine_question_and_answer_management",\
                        "rhine_spec_functional_requirements"): 
            issues = redmine.issue.filter(project_id=proj_id,assigned_to_id = 145)
                
            for issue in issues:
                try:

                    new_note =""
                    journal_list = issue.journals
                    for journal in journal_list:
                        if journal.notes != "":
                            new_note = journal.notes
                    found_row = 1000
                    if str(issue.id) in tickets_list:
                        found_row = int(tickets_list[str(issue.id)])
                        old_note = getCellValueinString(tickets_sheet.Cells(found_row,10))
                        if old_note != new_note:
                            tickets_sheet.Cells(found_row,10).Value = new_note
                            tickets_sheet.Cells(found_row,10).font.color = rgb_to_hex((255,0,255))
                        else:
                            tickets_sheet.Cells(found_row,10).font.color = rgb_to_hex((0,0,0))

                        del tickets_list[str(issue.id)]
                    
                    else:
                        tickets_sheet.Cells(wRow,10).Value = new_note
                        tickets_sheet.Cells(wRow,10).font.color = rgb_to_hex((255,0,255))
                        tickets_sheet.Cells(wRow,11).Value = proj_id
                        found_row = wRow
                        wRow += 1

                    if proj_id == "rhine":
                        tickets_sheet.Cells(found_row,1).Value = issue.id
                        tickets_sheet.Cells(found_row,2).Value = issue.subject 
                        tickets_sheet.Cells(found_row,5).Value = str(issue.due_date.year) + "/" + str(issue.due_date.month) + "/" + str(issue.due_date.day)
                        tickets_sheet.Cells(found_row,7).Value = issue.tracker.name
                        tickets_sheet.Cells(found_row,6).Value = issue.custom_fields[1].value
                        tickets_sheet.Cells(found_row,3).Value = issue.custom_fields[5].value
                        tickets_sheet.Cells(found_row,4).Value = "comment" + issue.custom_fields[6].value

                        work_time = 0
                        for time_entry in issue.time_entries:
                            print(time_entry)
                            if time_entry.user.id == 145:
                                work_time += float(time_entry.hours)
                                tickets_sheet.Cells(found_row,12).Value = str(work_time)


                except Exception as e:
                    print(e) 

        for remain_id in tickets_list:
            found_row = int(tickets_list[str(remain_id)])
            tickets_sheet.Cells(found_row,1).font.color = rgb_to_hex((178,178,178))                    

        tickets_wb.Close(SaveChanges = 1)
        return issues


    def getTeslinRedmineSubject(self):
        # hxy
        user_id = 145
        redmine = Redmine('http://iredmine.iauto.com/', key='d27046e8d2a28cee929411b668c25da4d47453b3', raise_attr_exception = False)

        #jinwan
        #redmine = Redmine('http://iredmine.iauto.com/', key='757f2035a324a88165543baf549c11f60d928936', raise_attr_exception = False) 
        #user_id = 164

        #sunhong
        #redmine = Redmine('http://iredmine.iauto.com/', key='582be232eabc6a850ef23d69d98117a1a618e792', raise_attr_exception = False) 
        #user_id = 367

    
        self.resultBook.ActiveSheet.Name = "Issue list"
        resultSheet = self.resultBook.ActiveSheet

        #for proj_id in ("rhine") : , "rhine_subject_management", "rhine_question_and_answer_management", "rhine_spec_functional_requirements"):
        proj_id = "rhine"
        issues = redmine.issue.filter(project_id=proj_id, status_id = "*", created_on ='>=2023-09-01') # ,  created_on='><2012-03-01|2012-03-07'
        wRow = 2
        for issue in issues:
            try:
                #issue = redmine.issue.get(issue_id)
                if issue.author.id == user_id:
                    resultSheet.Cells(wRow,1).Value = issue.id
                    resultSheet.Cells(wRow,2).Value = issue.subject
                    resultSheet.Cells(wRow,3).Value = str(issue.author.id)
                    resultSheet.Cells(wRow,5).Value = str(issue.created_on.year) + "/" + str(issue.created_on.month) + "/" + str(issue.created_on.day)
                    resultSheet.Cells(wRow,7).Value = issue.tracker.name
                    resultSheet.Cells(wRow,9).Value = proj_id
                    resultSheet.Cells(wRow,10).Value = str(issue.created_on.year) + "/" + str(issue.created_on.month) + "/" + str(issue.created_on.day)
                    wRow += 1
            
                journal_list = issue.journals
                for journal in journal_list:
                    if journal.notes != "":
                        if journal.user.id == user_id:
                            resultSheet.Cells(wRow,1).Value = issue.id
                            resultSheet.Cells(wRow,2).Value = issue.subject 
                            resultSheet.Cells(wRow,3).Value = str(journal.user.id)
                            resultSheet.Cells(wRow,5).Value = str(issue.created_on.year) + "/" + str(issue.created_on.month) + "/" + str(issue.created_on.day)
                            resultSheet.Cells(wRow,7).Value = issue.tracker.name
                            resultSheet.Cells(wRow,9).Value = proj_id
                            resultSheet.Cells(wRow,10).Value = str(journal.created_on.year) + "/" + str(journal.created_on.month) + "/" + str(journal.created_on.day)
                            resultSheet.Cells(wRow,11).Value = journal.notes
                            wRow += 1

                work_time = 0
                for time_entry in issue.time_entries:
                    print(time_entry)
                    if time_entry.user.id == user_id:
                        work_time = float(time_entry.hours)
                        resultSheet.Cells(wRow,1).Value = issue.id
                        resultSheet.Cells(wRow,2).Value = issue.subject 
                        resultSheet.Cells(wRow,3).Value = str(time_entry.user.id)
                        resultSheet.Cells(wRow,5).Value = str(issue.created_on.year) + "/" + str(issue.created_on.month) + "/" + str(issue.created_on.day)
                        resultSheet.Cells(wRow,7).Value = issue.tracker.name
                        resultSheet.Cells(wRow,9).Value = proj_id
                        resultSheet.Cells(wRow,12).Value = time_entry.id
                        resultSheet.Cells(wRow,13).Value = str(time_entry.updated_on.year) + "/" + str(time_entry.updated_on.month) + "/" + str(time_entry.updated_on.day)
                        resultSheet.Cells(wRow,14).Value = time_entry.comments
                        resultSheet.Cells(wRow,15).Value = str(time_entry.spent_on.year) + "/" + str(time_entry.spent_on.month) + "/" + str(time_entry.spent_on.day)
                        resultSheet.Cells(wRow,16).Value = str(work_time)
                        wRow += 1                    


            except Exception as e:
                print(e) 
        

        self.resultBook.Close(SaveChanges = 1)


    def updateTeslinRedmineTickets(self, due_date):
        tickets_file = os.getcwd() + "\\doc\\24MM_redmine_infor.xlsx"
        tickets_wb  = self.excelApp.Workbooks.Open(tickets_file)
        tickets_sheet = tickets_wb.Sheets('retreive') 

        #hxy
        redmine = Redmine('http://iredmine.iauto.com/', key='d27046e8d2a28cee929411b668c25da4d47453b3', raise_attr_exception = False)
        user_id = 145
        #sunhong
        #redmine = Redmine('http://iredmine.iauto.com/', key='582be232eabc6a850ef23d69d98117a1a618e792', raise_attr_exception = False) 
        #user_id = 367

        startRow = 2
        info = tickets_sheet.UsedRange
        nrows = info.Rows.Count
        tickets_list = {}
        while startRow <= nrows:
            if not isEmptyValue(tickets_sheet.Cells(startRow,1)):
                try:    
                    work_time = 0
                        
                    issue_id = getCellValueinString(tickets_sheet.Cells(startRow, 1))
                    time_comment = getCellValueinString(tickets_sheet.Cells(startRow, 14))
                    time_spend = getCellValueinString(tickets_sheet.Cells(startRow, 15))
                    time_add = getCellValueinString(tickets_sheet.Cells(startRow, 16))
                    issue = redmine.issue.get(int(issue_id))
                    work_time_add = 10000 # getCellValueinString(tickets_sheet.Cells(startRow, 13))
                    
                    for time_entry in issue.time_entries:
                        if time_entry.user.id == user_id:
                            work_time += float(time_entry.hours)
                    
                    tickets_sheet.Cells(startRow,20).value = str(work_time)
                    tickets_sheet.Cells(startRow,21).value = issue.user_id
                    tickets_sheet.Cells(startRow,22).value = str(issue.due_date)
                    
                    #补填工时time entry
                    if float(time_add) > 0:                
                        time_entry = redmine.time_entry.new()
                        time_entry.issue_id = issue_id
                        time_entry.spent_on = datetime.date(int(time_spend.split(",")[0]), int(time_spend.split(",")[1]),int(time_spend.split(",")[2]))
                        time_entry.hours = float(time_add) 
        
                        time_entry.activity_id = 10
                        time_entry.user_id = user_id
                        time_entry.comments = time_comment
                        time_entry.custom_fields=[{'id': 120, 'name':'工作包','value': '1001'}]
                        time_entry.save()
                    
                except Exception as e:
                    print(e)  
                    tickets_sheet.Cells(startRow,26).value = str(work_time)
                print(issue.id)
            startRow += 1

        '''
        for proj_id in ("rhine","rhine_client_input_management","rhine_subject_management", \
                        "rhine_question_and_answer_management", "rhine_spec_functional_requirements"): 
            try:
                issues = redmine.issue.filter(project_id=proj_id, assigned_to_id = 145)
                for issue in issues:
                    issue.due_date=due_date
                    print(issue.id)                    
                      
                    issue.save()
            except Exception as e:
                print(e) 
        '''

        return 0
    

    def getDocumentFromDNTC(self):
        try:
            print(sys.getfilesystemencoding())

            print(locale.getpreferredencoding())

            user_name = u'成牧(商泰)'
            my_confluence = Confluence(
                url = u'http://172.28.199.46/pages/viewpage.action?',
                username = user_name,
                password = u'shangtai334')

            label_title = "CCS5.0开发文档 商泰"
            
            label_title.encode('utf-8').decode("latin1")

            my_confluence.get_all_spaces(start=0,limit=500,expand=None)

            '''
            url_base = "view-source:http://172.28.199.46/spaces/listattachmentsforspace.action?key=CCS54"
            url_main = url_base + "pageId=25068012"
            html = requests.get(url_main, headers=header, timeout=timeout).text
            
            print(html)
            '''
        except Exception as e:
            print(e)        

    def getDataFromZiliaoZhan(self):

        try:
            url_base = 'http://down.ziliaozhan.org'
            url_main = url_base + "/books/pdf/神学/"
            html = requests.get(url_main, headers=header, timeout=timeout).text

            uk_mp3_file = None
            uk_mp3_url = None
            us_mp3_file = None
            us_mp3_url = None

            keyword = "&lt;dir&gt;"
            print("get data frome ziliaozhan")

            index = 0
            # Get entry feature
            while(index >= 0):
                index = html.find(keyword)
                if index >= 0 :
                    html = html[index+len(keyword):]
                    index = html.find("A HREF=")
                    if index >= 0:
                        html = html[index+len("A HREF="):]
                        index = html.find(">")
                        sub_url = html[1:index-2]
                        sub_url = url_base + sub_url
                        print(sub_url)
                        subhtml = requests.get(sub_url, headers=header, timeout=timeout).text

                        sub_index = subhtml.find(".pdf")
                        if sub_index >= 0:
                            title_index = subhtml.find("/</title>")
                            title_name = subhtml[0:title_index]
                            inx = title_name.rfind("/")
                            title_name = title_name[inx+1:]
                            subfolder = os.path.join("C:\\workspace\\temp\\神学", title_name)
                            if not os.path.exists(subfolder):
                                os.mkdir(subfolder)

                            sub_index = subhtml.find("[转到父目录]")
                            if sub_index >=0 :
                                subhtml = subhtml[sub_index:]

                            while(sub_index >=0):
                                
                                sub_index = subhtml.find("A HREF=")
                                if sub_index >= 0:
                                    subhtml = subhtml[sub_index + len("A HREF="):]
                                    sub_index = subhtml.find(">")
                                    pdf_url = subhtml[1:sub_index-1]
                                    pdf_url = url_base + pdf_url
                                    print(pdf_url)
                                    subhtml = subhtml[sub_index:]
                                    filename_index = subhtml.find("</A>")
                                    filename = subfolder + "\\" + subhtml[1: filename_index]

                                    res_length = requests.get(pdf_url, stream=True)
                                    total_size = int(res_length.headers['Content-Length'])
                                    print(res_length.headers)
                                    print(res_length)

                                    download_flag = False
                                    if os.path.exists(filename) and filename.find(".zip") < 0:
                                        temp_size = os.path.getsize(filename)
                                        print("当前：%d 字节， 总：%d 字节， 已下载：%2.2f%% " % (temp_size, total_size, 100 * temp_size / total_size))
                                        if temp_size >= total_size:
                                            download_flag = True
                                    else:
                                        temp_size = 0
                                        print("总：%d 字节，开始下载..." % (total_size,))

                                    if not download_flag:
                                        headers = {'Range': 'bytes=%d-' % temp_size,
                                                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:81.0) Gecko/20100101 Firefox/81.0"}
                                        res_left = requests.get(pdf_url, stream=True, headers=headers)

                                        with open(filename, "ab") as f:
                                            for chunk in res_left.iter_content(chunk_size=1024):
                                                temp_size += len(chunk)
                                                f.write(chunk)
                                                f.flush()

                                                done = int(50 * temp_size / total_size)
                                                sys.stdout.write("\r[%s%s] %d%%" % ('█' * done, ' ' * (50 - done), 100 * temp_size / total_size))
                                                sys.stdout.flush()

                                        #r = requests.get(pdf_url, allow_redirects=True)
                                        #open(filename, 'wb').write(r.content)

                        html = html[index:]
                else:
                    print("word feature not found")

        except Exception as e:
            print("error occur")

    def testRedmineApi(self):

        try:
            url_base = 'http://iredmine.storm'
            url_main = url_base + "/favorite_projects/search"
            response = requests.get(url_main, headers=header, timeout=timeout, auth=('huangxiaoyan', 'HackenLee00'))
            response.encoding = "utf-8"
            retreived_html = response.text

            retreived_url = response.url
            retreived_json = response.json

            root = ET.XML(html)
            
            print("get data frome ziliaozhan")

            
        except Exception as e:
            print(e)

    def testPolarion(self):

        funcfile = os.getcwd() + "\\tempResearch\\func_5_02_Handsfree.xlsx"
        funcfileWb  = self.excelApp.Workbooks.Open(funcfile)
        sheetCount = (funcfileWb.Worksheets.Count)

        func_table = {}

        for i in range(1, sheetCount + 1):
            funcSheet = funcfileWb.Worksheets(i)
            sheet_name = funcfileWb.Worksheets(i).Name
            if sheet_name.upper() not in ("COVER","CONTENTS","HISTORY"):
                info = funcSheet.UsedRange
                nrows = info.Rows.Count
                ncols = info.Columns.Count

                rowNo = 3
                startRow = 0
                endRow = 0
                description_str=""
                title_str = ""
                ncols = 36
                colNo = 3

                while colNo < ncols:
                    col_Value = getCellValueinString(funcSheet.Cells(4, colNo))
                    description_str = description_str + " " + col_Value
                    colNo += 1

                description_str = description_str + "\n"

                while colNo < ncols:
                    col_Value = getCellValueinString(funcSheet.Cells(6, colNo))
                    description_str = description_str + col_Value
                    colNo += 1


                func_table[sheet_name + "Heading"] = description_str
                while rowNo <= nrows:
                    if not isEmptyValue(funcSheet.Cells(rowNo, 1)) :
                        if title_str != "":
                            func_table[title_str] = description_str
                        title_str = getCellValueinString(funcSheet.Cells(rowNo, 1)) \
                                    + getCellValueinString(funcSheet.Cells(rowNo, 2)) \
                                    + getCellValueinString(funcSheet.Cells(rowNo, 3)) + " "
                        colNo = 3
                        while colNo < ncols:
                            title_str = title_str + getCellValueinString(funcSheet.Cells(rowNo, colNo))
                            colNo += 1
                        title_str = title_str.replace("Func.","")
                        description_str=""
                    else:
                        colNo = 1
                        while colNo < ncols:
                            col_Value = getCellValueinString(funcSheet.Cells(rowNo, colNo))
                            description_str = description_str + " " + col_Value
                            colNo += 1
                    rowNo += 1
                
                if title_str != "":
                    func_table[title_str] = description_str


        self.createWorkItemToPolarion(func_table,"https://svn.ci.iauto.com/svn/smartauto/Suzuki2024/03_Spec/func_5_01_Bluetooth.xlsx", "蓝牙")

        funcfileWb.Close(SaveChanges = 0)

    def statistic_21mm_page_number(self):
        rfqListfile = os.getcwd() + "\\doc\\24MM_statistic.xlsx"        
        rfqListWb  = self.excelApp.Workbooks.Open(rfqListfile)
        rfqListSheet = rfqListWb.Sheets('issues')

        wRow = 2
        info = rfqListSheet.UsedRange
        nrows = info.Rows.Count
        ''' 
        wRow = 724
        for root, dirs, files in os.walk("中国式样书3.00"):
            for name in files: 
                if name not in (".DS_Store","Thumbs.db"):
                    print("start extract this file:", name)
                    fileName = os.path.join(root, name)

                    rfqListSheet.Cells(wRow, 2).Value = fileName

                    try:
                        file = os.getcwd() + "\\" + fileName
                        if fileName.find(".doc") > 0:
                            rfqListSheet.Cells(wRow, 3).Value = self.getWordDocPageNumber(file)                            
                        elif fileName.find(".xl") > 0:
                            rfqListSheet.Cells(wRow, 3).Value = self.getExcelSheetNumber(file)
                        elif fileName.find(".ppt") > 0:
                            rfqListSheet.Cells(wRow, 3).Value = self.getPPTPageNumber(file)
                        elif fileName.find(".pdf") > 0:
                            rfqListSheet.Cells(wRow, 3).Value = self.getPDFPageNumber(file)
                        else:
                            rfqListSheet.Cells(wRow, 3).Value = 1
                    except Exception as e:
                        print("Open file error", fileName)
                        print(e)

                        rfqListSheet.Cells(wRow, 3).Value = "-"

                    wRow += 1

        rfq_filelist = {}
        while (wRow <724):
            rfq_filelist[getCellValueinString(rfqListSheet.Cells(wRow, 2))] = str(wRow)
            wRow += 1
        bFound = False

        '''     

        wRow = 2
        for root, dirs, files in os.walk("res\\HUAWEI_input"):
            for name in files:
                if name not in (".DS_Store","Thumbs.db"):
                    print("start extract this file:", name)
                    fileName = os.path.join(root, name)
                    rfqListSheet.Cells(wRow, 1).Value = fileName
                    if isEmptyValue(rfqListSheet.Cells(wRow, 3)):
                        try:
                            file = os.getcwd() + "\\" + fileName

                            if fileName.find(".doc") > 0:
                                rfqListSheet.Cells(wRow, 3).Value = self.getWordDocPageNumber(file)                            
                            elif fileName.find(".xl") > 0:
                                rfqListSheet.Cells(wRow, 3).Value = self.getExcelSheetNumber(file)
                            elif fileName.find(".ppt") > 0:
                                rfqListSheet.Cells(wRow, 3).Value = self.getPPTPageNumber(file)
                            elif fileName.find(".pdf") > 0:
                                rfqListSheet.Cells(wRow, 3).Value = self.getPDFPageNumber(file)
                            else:
                                rfqListSheet.Cells(wRow, 3).Value = '该文件类型无法统计'

                            wRow += 1
                        except Exception as e:
                            print("Open file error")
                            print(e)

        rfqListWb.Close(SaveChanges = 1)

    def getWordDocPageNumber(self, file):
        document = self.wordApp.Documents.Open(file)
        pageNumber = document.ComputeStatistics(2)
        document.Close()
        return pageNumber

    def getExcelSheetNumber(self, file):                 
        document  = self.excelApp.Workbooks.Open(file)
        sheetCount = document.Worksheets.Count                            
        document.Close(SaveChanges=0)

        return sheetCount
    
    def getPPTPageNumber(self, file):
        rfqPPT = self.pptApp.Presentations.Open(file)
        pageNumber = rfqPPT.slides.count
        rfqPPT.Close()

        return pageNumber
    
    def getPDFPageNumber(self, file):
        pdf = fitz.open(file)                       
        pageNumber = pdf.page_count
        pdf.close()

        return pageNumber

    def extractOutline_rfq(self):
        rfqListfile = os.getcwd() + "\\24MM_要件分析表_模板.xlsx"        
        rfqListWb  = self.excelApp.Workbooks.Open(rfqListfile)
        rfqListSheet = rfqListWb.Sheets('rfqlist')

        wRow = 3
        for root, dirs, files in os.walk("20211210 入手24MM LEXUS硬件RFQ"):
            for name in files: 
                if name not in (".DS_Store","Thumbs.db"):
                    print("start extract this file:", name)
                    fileName = os.path.join(root, name)

                    rfqListSheet.Cells(wRow, 5).Value = "RFQ"
                    last_folder = root.split("\\")
                    last_folder_name = last_folder[len(last_folder) - 1]
                    last_folder_name = last_folder_name.replace("●", "")
                    last_folder_name = last_folder_name.lstrip("_")
                    rfqListSheet.Cells(wRow, 5).Value = "RFQ"
                    rfqListSheet.Cells(wRow, 6).Value = last_folder_name
                    rfqListSheet.Cells(wRow, 7).Value = root
                    rfqListSheet.Cells(wRow, 8).Value = name
                    rfqListSheet.Cells(wRow, 11).Value = "2021/12/10"

                    try:
                        file = os.getcwd() + "\\" + fileName
                        outline_list = []

                        if fileName.find(".doc") > 0:
                            self.openWordDoc(file, outline_list)
                        elif fileName.find(".xl") > 0:
                            self.openExcel(file, outline_list)
                        elif fileName.find(".ppt") > 0:
                            self.openPPT(file, outline_list)

                        elif fileName.find(".pdf") > 0:
                            if name.upper().find("COVER") < 0:
                                #try:
                                
                                ##    in_file = os.path.abspath(file)
                                #    filename = file.split('\\')[-1]
                                #    out_file = os.path.abspath(os.getcwd() + '\\word\\' + filename[0:-4] + ".docx")
                                
                                #    result = self.convert_pdf2docx(in_file, out_file)
                                #except Exception as e:
                                #    print(e)
                                #    print(" Convert error")
                                
                                self.openPDF(file, outline_list)
                        else:
                            print("this file is not word/excel/ppt")

                        for chapter in outline_list:
                            rfqListSheet.Cells(wRow, 5).Value = "RFQ"
                            rfqListSheet.Cells(wRow, 6).Value = last_folder_name
                            rfqListSheet.Cells(wRow, 7).Value = root
                            rfqListSheet.Cells(wRow, 8).Value = name
                            rfqListSheet.Cells(wRow, 11).Value = "2021/12/10"
                            rfqListSheet.Cells(wRow, 10).Value = chapter

                            wRow += 1
                    except Exception as e:
                        print("Open file error")
                        print(e)
                        copyspecFile(fileName, "wrong file")

                    wRow += 1
                
        rfqListWb.Close(SaveChanges = 1)

    def ExtractWordFile(self):
        self.resultBook.ActiveSheet.Name = "wordlist"
        wRow = 1
        resultSheet = self.resultBook.ActiveSheet
        for root,dirs,files in os.walk("doc\\"):
            for name in files: 
                file0 = os.getcwd() + "\\doc\\" + name
                outline_list0 = {}
                self.openWordDoc_content(file0, outline_list0)

                resultSheet.Cells(wRow , 1).Value = name
                for  key in outline_list0 :
                    resultSheet.Cells(wRow , 2).Value = key
                    content_list = outline_list0[key].split("\r")
                    for content in content_list:
                        resultSheet.Cells(wRow , 3).Value = content
                        wRow += 1

        self.resultBook.Close(SaveChanges = 1)

                

    def diffWordFile(self):

        file0 = os.getcwd() + "\\doc\\24MM 多感覚連携詳細機能仕様書_302.docx"
        outline_list0 = {}
        self.openWordDoc_content(file0, outline_list0)

        file1 = os.getcwd() + "\\doc\\24MM 多感覚連携詳細機能仕様書_303.docx"
        outline_list1 = {}
        self.openWordDoc_content(file1, outline_list1)

        self.resultBook.ActiveSheet.Name = "diff result"
        resultSheet = self.resultBook.ActiveSheet

        for key in outline_list0:
            if key in outline_list1:
                if outline_list0[key] == outline_list1[key]:
                    outline_list0[key] = "DEL"
                    outline_list1[key] = "DEL"

        row  = 1
        for key in outline_list0:
            if outline_list0[key] != "DEL":
                resultSheet.Cells(row, 1).Value = key
                resultSheet.Cells(row, 2).Value = outline_list0[key]
                row += 1

        for key in outline_list1:
            if outline_list1[key] != "DEL":
                resultSheet.Cells(row, 1).Value = key
                resultSheet.Cells(row, 2).Value = outline_list1[key]
                row += 1

        self.resultBook.Close(SaveChange = 1)

    def openWordDoc_content(self, file, outline_list):
        document = self.wordApp.Documents.Open(file)
        func_table = {}
        title_str = ""
        description_str = ""
        bStartRecord = False
        paragraphs = document.paragraphs 

        heading1_start = 0
        heading2_start = 0
        heading3_start = 0

        text = ""
        content_text = ""
        key_value = ""
        for i in range(0, paragraphs.count):
            para = paragraphs[i]
            text = para.Range.text
            style_name = para.style.NameLocal 
            print("word: " + style_name + " " + text)
                      
            if style_name.find("标题") >=0 or style_name.find("列出段落") >=0 :
                if key_value != "":
                    outline_list[key_value] = content_text
                key_value = (para.Range.Listformat.ListString + text).replace(" ","").replace("\r","")
                content_text = ""

            else:
                content_text += text

        document.Close()

    def openTMECWordDoc_content(self, file, outline_list):
        document = self.wordApp.Documents.Open(file)
        bStartRecord = False
        paragraphs = document.paragraphs 

        text = ""
        content_text = ""
        key_value = ""
        for i in range(0, paragraphs.count):
            para = paragraphs[i]
            text = para.Range.text
            style_name = para.style.NameLocal 
            print("word: " + style_name + " " + text)
            
                      
            if style_name.find("标题") >=0 or style_name.find("列出段落") >=0 :
                if key_value != "":
                    outline_list[key_value] = content_text.replace(" ","").replace("\r","")
                key_value = (para.Range.Listformat.ListString + text).replace(" ","").replace("\r","")
                content_text = ""

            else:
                content_text += text

    
    def openWordDoc(self, file, outline_list):
        document = self.wordApp.Documents.Open(file)
        func_table = {}
        title_str = ""
        description_str = ""
        bStartRecord = False
        paragraphs = document.paragraphs 

        heading1_start = 0
        heading2_start = 0
        heading3_start = 0

        isContent = False
        text = ""
        content_text = ""
        for i in range(0, paragraphs.count):
            para = paragraphs[i]
            text = para.Range.text
            style_name = para.style.NameLocal 
            print("word: " + style_name + " " + text)
            
            
            if style_name.startswith("Heading") or style_name == "1" or style_name == "2" or style_name == "3" :
                number_arrange = ""
                if style_name == "1":
                    heading1_start += 1
                    number_arrange = str(heading1_start) + "."
                    heading2_start = 0
                    heading3_start = 0
                elif style_name == "2":
                    heading2_start += 1
                    number_arrange = str(heading1_start) + "." + str(heading2_start)
                    heading3_start = 0
                elif style_name == "3":
                    heading3_start += 1
                    number_arrange = str(heading1_start) + "." + str(heading2_start)+ "." + str(heading3_start)

                
                if text.find(str(heading1_start) + ".") < 0 :
                    text = number_arrange + " " + text
            
                outline_list.append(text)

            if style_name.startswith("目录"):
                text = text[0:text.rfind("\t")]
                outline_list.append(text)
                isContent = True
            elif style_name == "1" or style_name == "2" or style_name == "3" or isContent == True:
                break
            
            if para.Range.Style.Font.Strikethrough == 0:
                content_text += text

        if isContent == False:
            content_text = content_text.split("\r")
            for content in content_text:
                content = content.replace(" ","")
                if  content != "":
                    outline_list.append(content)        

        document.Close()

    def openExcel(self, file, outline_list):
        rfqTableWb  = self.excelApp.Workbooks.Open(file)
        
        sheetCount = rfqTableWb.Worksheets.Count
        for i in range(1, sheetCount + 1):
            sheet_name = rfqTableWb.Worksheets(i).Name
            outline_list.append(sheet_name)

        rfqTableWb.Close(SaveChanges = 0)
    
    def openPPT(self, file, outline_list):
        rfqPPT = self.pptApp.Presentations.Open(file)
        rfqslides = rfqPPT.slides

        isContentPage = False
        content_text = ""        
        for oneSlide in rfqslides:
            #if isContentPage:
            #    break
            shapes = oneSlide.shapes

            for shape in shapes:
                print ("ppt : ", shape.name)
                text = ""
                try:
                    text = shape.textframe.TextRange.text
                except Exception as e:
                    text = ""
                if text.replace(" ","").find("目录") >= 0 or \
                    text.replace(" ","").find("目次") >= 0 or \
                    text.replace(" ","").upper().find("CONTENT") >= 0:
                    isContentPage = True
                else:
                    text = text.replace("\u3000", "")
                    print(text)
                    
                content_text += text

        content_text_list = content_text.split("\r")
        for content in content_text_list:
            if content != "":
                outline_list.append(content)

        rfqPPT.Close()

    def openPDF(self, file, outline_list):
        pdf = fitz.open(file)

        isContent = False
        content_text = ""
        whole_text = ""
        chk_page_number = int(pdf.page_count * 0.1)
        for page in pdf:
            page_text = page.get_text()
            print("PDF: ", page_text)
            pure_text = page_text.replace(" ","")
            if pure_text.find("目次") >= 0 or pure_text.upper().find("CONTENT") >= 0:
                isContent = True
                chk_page_number = page.number + 4
            #elif isContent == True and page.number > chk_page_number:
            #    break

            if page.firstLink is not None or isContent:
                content_text += "page number " + str(page.number) + ":" + page_text

            whole_text += "page number " + str(page.number) + ":" + page_text

        '''
        if isContent :
            content_text = content_text.replace("TABLE OF CONTENTS／目次","")
            content_text = content_text.replace("目 次","")
            content_text = content_text.replace("目次","")
            content_text = content_text.replace("TOYOTA MOTOR CORPORATION","")
        '''

        content_text = content_text.split("\n")
        for content in content_text:
            if content.replace(" ","") != "":
                content = content[0:content.find("..")]
                outline_list.append(content)
        
        if len(outline_list) == 0:
            whole_text = whole_text.split("\n")
            for whole in whole_text:
                if whole.replace(" ","") != "":
                    whole = whole[0:whole.find("..")]
                    outline_list.append(whole)


        '''
        with open(file, 'rb') as f:
            pdf = PdfFileReader(file)
            if not pdf.isEncrypted:
                outline = pdf.getOutlines()
                print(outline)
        
                print("=====extract text======")
                for i in range(0, pdf.numPages):
                    print("page :", str(i))
                    pageObj = pdf.getPage(i)
                    print(pageObj.extractText())
                    #.encode("uft8")
            else:
                print("This pdf is encrypted")        
        '''

        pdf.close()

    def chkRfq(self):
        original_file = os.getcwd() + "\\24MM_要件分析表_模板.xlsx"
        originalWb  = self.excelApp.Workbooks.Open(original_file)       
        amend_file =  os.getcwd() + "\\24MM_要件分析表_第一次导入rfq_211201.xlsx"
        amendWb  = self.excelApp.Workbooks.Open(amend_file)       
        amendSheet = amendWb.Sheets("Sheet1")

        orignalSheet = originalWb.Sheets('Sheet1')
        info = orignalSheet.UsedRange
        nrows = info.Rows.Count

        original_content_list = []

        row = 3
        while(row <= nrows):
            original_content_list.append((getCellValueinString(orignalSheet.Cells(row, 8)) + getCellValueinString(orignalSheet.Cells(row, 10))).replace(" ","").upper())
            row += 1
        
        info = amendSheet.UsedRange
        nrows = info.Rows.Count

        row = 3
        while(row <= nrows):
            amendContent = (getCellValueinString(amendSheet.Cells(row, 8)) + getCellValueinString(amendSheet.Cells(row, 10))).replace(" ","").upper()
            if amendContent in original_content_list:
                amendSheet.Cells(row, 1).Value = "CHK"
            row += 1

        amendWb.Close(SaveChanges = 1)
        originalWb.Close(SaveChanges = 0)

    def testWordApp(self):

        funcfile = os.getcwd() + '\\doc\\A88_Func_18_本地音乐.docx'
        document = Document(funcfile)
        func_table = {}
        title_str = ""
        description_str = ""
        bStartRecord = False
        paragraphs = document.paragraphs 


        root_node = None
        last_node = None
    
        for i in range(0, len(paragraphs)):
            style_name = document.paragraphs[i].style.name
            text = document.paragraphs[i].text
            print(style_name)

            if style_name.startswith("Heading"):
                new_node = nodeObj()
                style_level = int(style_name.replace("Heading ",""))

                new_node.set_value(text)
                new_node.set_level(style_level)
                new_node.set_parent(None)
                new_node.set_child(None)
                new_node.set_sibling(None)
                if last_node is None:
                    root_node = new_node
                    last_node = root_node
                else:
                    if last_node.node_level < style_level:
                        last_node.set_child(new_node)
                        new_node.set_parent(last_node)
                    elif last_node.node_level == style_level:
                        last_node.set_sibling(new_node)
                        new_node.set_parent(last_node.get_parent())
                    else:
                        last_parent = last_node.get_parent()
                        last_parent.set_sibling(new_node)
                        new_node.set_parent(last_parent.get_parent())
                    
                    last_node = new_node

                description_str = ""
            else:
                if ((style_name.find("text") >= 0) or (style_name.find("Normal") >= 0)):
                    description_str = description_str + "<br/>\n"+ text
                    if last_node:
                        last_node.set_content(description_str)


            if text.find("功能要求") >= 0 and style_name.find("Heading 1") >= 0:
                bStartRecord = True
                title_str = "本地音乐_" + text # + "|" + style_name
            else:
                if bStartRecord:
                    if style_name.find("Heading") >= 0 and text.startswith("3."):
                        title_str = text    #+ "|" + style_name
                        description_str = ""

                    if ((style_name.find("text") >= 0) or (style_name.find("Normal") >= 0)):
                        description_str = description_str + "<br/>\n"+ text

                    if title_str != "":
                        func_table[title_str] = description_str

            i += 1

        external_url = "https://svn.ci.iauto.com/svn/smartauto/A88_Func_18_本地音乐.docx"
        func_name = "本地音乐"

        #self.createWorkItemToPolarion(func_table, external_url, func_name)
        self.createWorkItemToPolarionFromDomTree(root_node, external_url, func_name)

    def createWorkItemToPolarionFromDomTree(self, root_node, url, func_name):
        #client = polarion.Polarion('https://almdemo.polarion.com/polarion', 'georgesandlily@gmail.com', 'jN!P6zRXs')
        client = polarion.Polarion('https://almdemo.polarion.com/polarion', 'huangxiaoyan@iauto.com', 'x!NM@5wJ*')
        project = client.getProject('testproj')

        first_workitem = None
        parent_workitem = None

        current_node = root_node

        while current_node is not None:

            title_str = current_node.node_value
            if title_str.startswith("3"):

                new_workitem = project.createWorkitem('systemrequirement')
                description_str = current_node.node_content
                new_workitem.setDescription(description_str)

                new_workitem.save()

                current_node.set_extraVal(new_workitem.id)

                reload_workitem = project.getWorkitem(new_workitem.id)
                reload_workitem.addComment("机能分类", func_name)
                reload_workitem.addHyperlink(url,  reload_workitem.HyperlinkRoles.EXTERNAL_REF)
            
                reload_workitem.title = title_str
                reload_workitem.save()
                print("create a new workitem:", new_workitem.id)

                if current_node.get_parent() is not None:
                    reload_workitem = project.getWorkitem(new_workitem.id)
                    item_id = current_node.get_parent().node_extraVal
                    linked_item = project.getWorkitem(item_id)
                    reload_workitem.addLinkedItem(linked_item, 'has parent')
                    print("---- ", title_str + " has parent " + linked_item.title)                

                    #reload_workitem.save()

            next_node = current_node.get_child()
            if next_node is None:
                next_node = current_node.get_sibling()
                if next_node is None:
                    next_node = current_node.get_parent().get_sibling()
                    
            current_node = next_node
             
    def createWorkItemToPolarion(self, func_table, url, func_name):
        '''
        client = polarion.Polarion('https://almdemo.polarion.com/polarion', 'huangxiaoyan@iauto.com', 'x!NM@5wJ*')
        project = client.getProject('Lexus_test')
        
        #workitem = project.getWorkitem("TEST-722")       
        #workitem.customFields.Custom.append("Function A")
        #workitem.setCustomField('Function A', func_name)
        
        first_workitem = None
        parent_workitem = None
        for key in func_table:

            try:
                new_workitem = project.createWorkitem('systemrequirement')
                description_str = func_table[key]
                new_workitem.setDescription(description_str)

                new_workitem.save()

                reload_workitem = project.getWorkitem(new_workitem.id)
                reload_workitem.addComment("机能分类", func_name)
                reload_workitem.addHyperlink(url,  reload_workitem.HyperlinkRoles.EXTERNAL_REF)
                
                if key.find("Heading") >= 0 :
                    key = key.replace("Heading", "")
                    parent_workitem = reload_workitem
                else:
                    reload_workitem.addLinkedItem(parent_workitem, "has parent")

                reload_workitem.title = key
                reload_workitem.save()

                print("create a new workitem:", new_workitem.id)
            except Exception as e:
                print(e)
        '''
        print("")

    def importExcelToPolarion(self):
        client = polarion.Polarion('https://almdemo.polarion.com/polarion', 'huangxiaoyan@iauto.com', 'x!NM@5wJ*')
        project = client.getProject('Lexus_test')

        spaces = project.getDocumentSpaces()
        #document = project.createDocument(spaces[1], 'VideoPlayer模块软件需求分析说明书','VideoPlayer模块软件需求分析说明书',  ['requirement'], 'relates_to')
        key = spaces[1] + "//VideoPlayer模块软件需求分析说明书"
        document = project.getDocument(key)
        reload_workitem = project.getWorkitem('LEXU-577')       
        
        document.addHeading('VideoPlayer模块软件需求分析说明书')
        chapter_1 = document.addHeading('封面')
        
        document.addHeading('项目编号：Project Name', chapter_1)
        document.addHeading('模块版本：V1.0', chapter_1)

        
        chapter_2 = document.addHeading('1.功能统括')
        chapter_2_1 = document.addHeading('1.1	功能概要', chapter_2)
        document.addHeading('本文档主要描述videoplayer功能', chapter_2_1)

        chapter_2_2 = document.addHeading('1.2	功能清单', chapter_2)
        document.addHeading('系统需求ID		功能ID		功能(L1)		功能(L2)		功能(L...)		L1		L2		T-EMV', chapter_2_2)	
        document.addHeading('REQ_01		ID01		多媒体选择		Audio		Video		○		○		×	', chapter_2_2)	

        chapter_2_3 = document.addHeading('1.3	缩略语', chapter_2)
        document.addHeading('编号 用语、缩略语 含义、定义和正式名称', chapter_2_3)
        chapter_3 = document.addHeading('2.功能要求')
        document.addHeading('功能ID	功能名	功能生效条 功能描述	需求可行性	优先级	依赖关系 验证准则 风险 优先级 依赖关系 前提条件	验证步骤 期望结果', chapter_3)	
        document.addHeading('ID01	媒体源选择	USB设备插入	"1.两个USB口对应两个媒体源\n 2.媒体源的选择与本地音乐的USB1和USB2选择相同"	可	高	无	-	"1. 两个USB口都不插U盘，进入视频播放器APP \
                            2. USB1口插入U盘\
                            3. USB2口插入U盘\
                            4. 拔出USB2口U盘"	"1. 进入APP状态栏USB1/USB2图标不可点击，画面显示请插入USB；\
                            2. USB1口插入U盘，Toast提示USB1已插入；\
                            3. USB2口插入U盘，Toast提示USB2已插入；\
                            4. 拔出USB2口U盘，Toast提示USB2已拔出；"	无', chapter_3)	
        										
        chapter_4 = document.addHeading('3.非功能要求')
        chapter_5 = document.addHeading('4.SDK和API功能要求')
        chapter_6 = document.addHeading('5.操作环境')


        #document.addHeading(page1_content, chapter_1)

        #new_workitem = project.createWorkitem('systemrequirement')
        #description_str = func_table[key]
        #new_workitem.setDescription(description_str)

        #new_workitem.save()

    def extract_pdf(self, root):
        self.excelApp.Workbooks.Add()
        resultBook = self.excelApp.ActiveWorkBook
        resultBook.ActiveSheet.Name = "pdf content"
        resultSheet =resultBook.ActiveSheet

        wRow = 1
        for root, dirs,files in os.walk(root):
            for name in files: 
                fileName = os.path.join(root, name)
                file = os.getcwd() + "\\" + fileName
                print("start analyz:", fileName)
                folder_name = name.replace(".pdf","")
                try:
                    wRow = self.grabImage_text(file, folder_name, resultSheet, wRow)
                except Exception as e:
                    print(e)

                #with open(file, 'rb') as f:
                #    pdf = PdfFileReader(f)
                #    if not pdf.isEncrypted:
                #        outline = pdf.getOutlines()
                #        print(outline)
                #        print("=====extract text======")
                #        resultSheet.Cells(wRow, 1).Value = fileName
                #        for i in range(0, pdf.numPages):
                #            print("page :", str(i))
                #            pageObj = pdf.getPage(i)
                #            print(pageObj.extractText())
                #            resultSheet.Cells(wRow, 2).Value = str(i)
                #            resultSheet.Cells(wRow, 3).Value = pageObj.extractText()
                #            wRow += 1
                #    else:
                #        print("This pdf is encrypted")

                #for page_layout in extract_pages(file):
                #    for element in page_layout:
                #        print(element)

                #try:
                #
                #    in_file = os.path.abspath(file)
                #    filename = file.split('\\')[-1]
                #    out_file = os.path.abspath(os.getcwd() + '\\word\\' + filename[0:-4] + ".docx")
                #
                #    result = self.convert_pdf2docx(in_file, out_file)
                #except Exception as e:
                #    print(e)
                #    print(" Convert error")
        
        resultBook.Close(SaveChanges = 1)
        return

    def convert_pdf2docx(self, input_file, output_file, pages: Tuple = None):
        
        #"Converts pdf to docx"
        if pages:
            pages = [int(i) for i in list(pages) if i.isnumeric()]

        result = parse(input_file, output_file, pages)
        summary = {
            "File": input_file, "Pages": str(pages), "Output File": output_file
        }

        # Printing Summary
        print("## Summary ########################################################")
        print("\n".join("{}:{}".format(i, j) for i, j in summary.items()))
        print("###################################################################")
        
        return result

    def grabImage_text(self, pdfPath, folder_name, resultSheet, wRow):
        pdf_img_path = os.path.join(os.getcwd(), folder_name) # + "\\imageGrab\\"
        if not os.path.exists(pdf_img_path):
            os.mkdir(pdf_img_path)

        self.convert_pdf2image(pdfPath, pdf_img_path, 1,1,0)

        for root, dirs,files in os.walk(pdf_img_path):
            for name in files: 
                try:
                    imageFile = pdf_img_path + "\\" + name
                    Image.MAX_IMAGE_PIXELS = None
                    text = pytesseract.image_to_string(Image.open(imageFile), lang='jpn+eng') #  chi_sim
                    textlines = text.split("\n")
                    for text_line in textlines:
                        if text_line != "":
                            resultSheet.Cells(wRow, 1).Value = folder_name
                            resultSheet.Cells(wRow, 2).Value = name
                            resultSheet.Cells(wRow, 3).Value = text_line
                            wRow += 1      
                except Exception as e:
                    print(e)
        return wRow

    def convert_pdf2image(self, pdfPath, imgPath,zoom_x,zoom_y,rotation_angle):
        pdf_Path = os.path.join(os.getcwd() + "\\" + pdfPath)
        
        img_Path = os.path.join(os.getcwd() + "\\"+ imgPath)
        # 打开PDF文件
        pdf = fitz.open(pdfPath)
        if not pdf.authenticate("CCS5.0HMI"):
            print('cannot decrypt %s with password %s' % (pdfPath, "CCS5.0HMI"))
            
        # 逐页读取PDF
        for pg in range(0, pdf.pageCount):
            try:
                page = pdf[pg]
                # 设置缩放和旋转系数
                trans = fitz.Matrix(zoom_x, zoom_y).prerotate(rotation_angle)
                pm = page.get_pixmap(matrix=trans, alpha=False)
                # 开始写图像
                pm.save(imgPath+"\\"+str(pg+1)+".png")
            except Exception as e:
                print(e)
        pdf.close()

    def upd_funcID_FORD(self):
        featurefile = os.getcwd() + "\\res\\Ford\\Ford China 【SYNC+4.0】CX771 & CX821 IVI FeatureList.xlsx"        
        featureWb  = self.excelApp.Workbooks.Open(featurefile)
        featureSheet = featureWb.Sheets('FeatureList')

        func_id_list = {}
        info = featureSheet.UsedRange
        nrows = info.Rows.Count
        row = 3
        
        while row <= nrows:
            key = getCellValueinString(funcSheet.Cells(row, 3)).upper()
            func_id_list[key] = getCellValueinString(funcSheet.Cells(row, 4))
            row += 1

        info = featureSheet.UsedRange
        nrows = info.Rows.Count
        row = 3
        
        while row <= nrows:
            if not isEmptyValue(featureSheet.Cells(row, 27)):
                func_id = getCellValueinString(featureSheet.Cells(row, 27)).upper()
                if func_id in func_id_list:
                    featureSheet.Cells(row, 28).Value = func_id_list[func_id]
            row += 1

        featureWb.Close(SaveChanges = 1)

    def add_schedule_FORD(self):
        schedulefile = os.getcwd() + "\\res\\Ford\\FordSYNC+4.0项目开发时间计划v1.0.xlsx"        
        scheduleWb  = self.excelApp.Workbooks.Open(schedulefile)
        scheduleSheet = scheduleWb.Sheets('Timing Plan')
        
        featurefile = os.getcwd() + "\\res\\Ford\\Ford China 【SYNC+4.0】CX771 & CX821 IVI FeatureList.xlsx"        
        featureWb  = self.excelApp.Workbooks.Open(featurefile)
        featureSheet = featureWb.Sheets('FeatureList')

        milestone_list = {}
        info = scheduleSheet.UsedRange
        nrows = info.Rows.Count
        row = 7
        while (row <= nrows+1):
            if not isEmptyValue(scheduleSheet.Cells(row, 4)):
                feature_key = getCellValueinString(scheduleSheet.Cells(row, 3)).upper().replace(" ","")
                milestone_list[feature_key] = getCellValueinString(scheduleSheet.Cells(row, 4))
            row += 1

        info = featureSheet.UsedRange
        nrows = info.Rows.Count
        row = 3
        
        while (row <= nrows+1):
            key = getCellValueinString(featureSheet.Cells(row, 5)).upper().replace(" ","")
            if key in milestone_list:
                featureSheet.Cells(row,20).Value = milestone_list[key]
                del milestone_list[key]
            row += 1

        wRow = 140
        for remain in milestone_list:
            featureSheet.Cells(wRow, 5).Value = remain            
            featureSheet.Cells(wRow, 20).Value = milestone_list[remain]
            wRow += 1

        featureWb.Close(SaveChanges = 1)

    def add_rfq_FORD(self):
        rfqListfile = os.getcwd() + "\\res\\Ford\\Ford China 【SYNC+4.0】CX771 & CX821 IVI traceability.xlsx"        
        rfqListWb  = self.excelApp.Workbooks.Open(rfqListfile)
        rfqListSheet = rfqListWb.Sheets('featureid')
        specfileSheet = rfqListWb.Sheets('filelist')
        feature_id_list = {}
        info = rfqListSheet.UsedRange
        nrows = info.Rows.Count
        row = 2
        while (row <= nrows):
            feature_id_list[getCellValueinString(rfqListSheet.Cells(row, 1))] = row
            row += 1

        info = specfileSheet.UsedRange
        nrows = info.Rows.Count
        row = 2
        
        while (row <= nrows):
            feature_id = getCellValueinString(specfileSheet.Cells(row, 1))
            if feature_id in feature_id_list:
                specfileSheet.Cells(row, 2).Value = getCellValueinString(rfqListSheet.Cells(feature_id_list[feature_id], 2))
            row += 1

        rfqListWb.Close(SaveChanges = 1)

    def copyWordfile_FORD(self, root):
        ford_word_list = [
                "FD_Comprehensive Status_V3.0_DRAFT.docx",
                "Digital Scent PRD - for CDX707C 20210714.doc",
                "Enhanced Memory InterfaceClient SPSS v1.13 July 30, 2021.docx",
                "Enhanced_Memory_With_Soft_Button_Feature_Level_Specification_v1.3_May_10, 2022.docx",
                "Lincoln Embrace_v2.2.2_CX821_Draft_V2.docx",
                "Sequential Animation IVI&Ambient Light Feature Spec (if equipped).docx",
                "Settings in Centerstack Logical to Physical CAN signal mapping Feb 5, 2019.docx",
                "WiFi Hotspot Server v2 SPSS v1.5 September 22, 2020.docx",
                "SYNC+3.0_Carrier_Subscription_Feature_Document_v1.0.docx",
                "【福特Phase4】MRD_低电量提醒_CX727_v1.4_20220602.docx",
                "Message Center Feature document --draft 1.0.docx",
                "PRD - Message Center_Product Team_Yao Jiang_20220531_V0.1.docx",
                "FD_Relax_v0.4.docx",
                "PRD - Wallbox-vehicle Communication_ECDX Charging_JZENG21_20220609_V1.0.docx",
                "Customer Feedback Loop- Sync+ Voice Message  v1.1-20210623(1).docx",
                "福特_PRD_随心唱_V1.3_20211008.docx",
                "PRD - Car Butler_Product Team_Yao Jiang_20220526_V0.1.docx",
                "Audio Man logical to physical Aug 26th.docx",
                "DUEROS_CVPP_Ford and Baidu Accoun_v2.0.7.docx",
                "CarPlay SPSS v1.3 Jun 2, 2017.docx",
                "FNV2 Climate Control System HMI Requirements (EFP_RCCM) vZ.docx",
                "Bluetooth_Security_Spec_1.0.docx",
                "C.13 Ford_CyberAssurance-SOW_ReleaseV1.3(EE).doc",
                "IVI Security Requirements v7.docx",
                "Key Material Requirements.docx",
                "Operating_Systems_Security_Requirements_RC3.docx",
                "S26_Secure_Channel_v1_04_Release.doc",
                "FORD Input Specification.docx",
                "IP Pass Through Client SPSS v1.1 September 4 2018 (002)_DNS.docx",
                "Power Management Variant 3 APIM_AOS SPSS v1.4 Feb 3, 2022.docx",
                "Wireless Interface Router Server SPSS v1.8 January 18, 2022.docx",
                "548171_H_001_IVSU_Vehicle_Function_HMI v4.0.docx",
                "ProjectH_FeatureDucumentV02.docx",
                "Selectable Drive Mode (SDM) Feature Specification_Gen 3_20210319.docm"]

        for root, dirs, files in os.walk(root):
            for name in files: 
                if name not in (".DS_Store","Thumbs.db"):
                    print("start extract this file:", name)
                    fileName = os.path.join(root, name)
                    if name in ford_word_list:
                        copyspecFile(fileName, "worddoc")

    def extractOutline_rfq_FORD(self,root):
        rfqListfile = os.getcwd() + "\\res\\pdf_outline.xlsx"        
        rfqListWb  = self.excelApp.Workbooks.Open(rfqListfile)
        #rfqListSheet = rfqListWb.Sheets('FeatureList')
        specfileSheet = rfqListWb.Sheets('RFQList')
        specfiles = {}
        info = specfileSheet.UsedRange
        nrows = info.Rows.Count
        row = 3
        wRow = 2
        '''
        while (row <= nrows):
            specfile_list = getCellValueinString(rfqListSheet.Cells(row, 10)).split("\n")
            for specfile in specfile_list:
                speckey = specfile.replace(".xlsx","").replace(".pptx","").replace(".ppx","").replace(".docx","").replace(".doc","").replace(".xls","").replace(".pdf","")
                if speckey in specfiles:
                    current_feature = specfiles[speckey]
                    specfiles[speckey] = current_feature + "@" + getCellValueinString(rfqListSheet.Cells(row, 4)) +  "|" + getCellValueinString(rfqListSheet.Cells(row, 5))
                else:
                    specfiles[speckey] = getCellValueinString(rfqListSheet.Cells(row, 4)) + "|" + getCellValueinString(rfqListSheet.Cells(row, 5))

            row += 1
        
        wRow = 3
   
        while (row <= nrows):
            specfile_list = getCellValueinString(rfqListSheet.Cells(row, 10)).split("\n")
            for specfile in specfile_list:
                speckey = specfile.replace(".xlsx","").replace(".pptx","").replace(".ppx","").replace(".docx","").replace(".doc","").replace(".xls","").replace(".pdf","")
                if speckey in specfiles:
                    current_feature = specfiles[speckey]
                    specfiles[speckey] = current_feature + "@" + getCellValueinString(rfqListSheet.Cells(row, 2))
                else:
                    specfiles[speckey] = getCellValueinString(rfqListSheet.Cells(row, 2))

            row += 1

        info = specfileSheet.UsedRange
        nrows = info.Rows.Count
        row = 2
        while(row <= nrows):
            name = getCellValueinString(specfileSheet.Cells(row, 1))
            namekey = name.replace(".xlsx","").replace(".pptx","").replace(".ppx","").replace(".docx","").replace(".doc","").replace(".xls","").replace(".pdf","")
            if namekey in specfiles:
                specfileSheet.Cells(row, 5).Value = specfiles[namekey]

            row += 1
        '''

        for root, dirs, files in os.walk(root):
            for name in files: 
                if name not in (".DS_Store","Thumbs.db"):
                    print("start extract this file:", name)
                    fileName = os.path.join(root, name)
                    namekey = name.replace(".xlsx","").replace(".pptx","").replace(".ppx","").replace(".docx","").replace(".doc","").replace(".xls","").replace(".pdf","")
                       
                    try:
                        file = os.getcwd() + "\\" + fileName
                        outline_list = []

                        if fileName.find(".doc") > 0:
                            self.openWordDoc(file, outline_list)
                        elif fileName.find(".xl") > 0:
                            self.openExcel(file, outline_list)
                        elif fileName.find(".ppt") > 0:
                            self.openPPT(file, outline_list)

                        elif fileName.find(".pdf") > 0:
                            if name.upper().find("COVER") < 0:

                                self.openPDF(file, outline_list)
                        else:
                            print("this file is not word/excel/ppt")

                        for chapter in outline_list:
                            specfileSheet.Cells(wRow, 1).Value = ""
                            specfileSheet.Cells(wRow, 2).Value = ""
                            specfileSheet.Cells(wRow, 3).Value = root
                            specfileSheet.Cells(wRow, 4).Value = name
                            specfileSheet.Cells(wRow, 5).Value = "RFQ"
                            specfileSheet.Cells(wRow, 7).Value = "2024/01/25"
                            specfileSheet.Cells(wRow, 6).Value = chapter

                            wRow += 1
                    except Exception as e:
                        print("Open file error")
                        print(e)
                    
                    copyspecFile(fileName, "ford_spec")

                    #del specfiles[namekey]
                    wRow += 1
        
        for remain in specfiles:  
            feature_infor = specfiles[remain].split("@")
            feature_id = ""                        
            feature_name = ""
            for feature in feature_infor:
                feature_id = feature_id + "\n" + feature.split("|")[0]
                feature_name = feature_name + "\n" + feature.split("|")[1]
            specfileSheet.Cells(wRow, 1).Value = feature_id
            specfileSheet.Cells(wRow, 2).Value = feature_name      
            specfileSheet.Cells(wRow, 4).Value = remain 
            wRow += 1
        
        rfqListWb.Close(SaveChanges = 1)
        

if __name__ == '__main__':
    print("Start running ", time.localtime(time.time()))
    metaF = ToolsFixer()
    if len(sys.argv) == 2:
        if sys.argv[1] == "GET_ZILIAO":
            metaF.getDataFromZiliaoZhan()
        elif sys.argv[1] == "TEST_REDMINE_TASK":
            metaF.createRedmineTask()
        elif sys.argv[1] == "TEST_REDMINE_UPDATE":
            metaF.updateRedmineStatus()
        elif sys.argv[1] == "TEST_REDMINE_SUBJECT":
            metaF.createNormalSpecChangeTickets()
        elif sys.argv[1] == "TEST_POLARION":
            metaF.testPolarion()
        elif sys.argv[1] == "TEST_WORD":
            metaF.testWordApp()
        elif sys.argv[1] == "TEST_IMAGEGRABTEXT":
            metaF.grabImage_text()
        elif sys.argv[1] == "TEST_EXTRACTRFQ":
            metaF.extractOutline_rfq()
        elif sys.argv[1] == "CHK_EXTRACTRFQ":
            metaF.chkRfq()
        elif sys.argv[1] == "DOWNLOAD_FROM_DNTC":
            metaF.getDocumentFromDNTC()
        elif sys.argv[1] == "STATISTIC_PAGE_NUMBER":
            metaF.statistic_21mm_page_number()
        elif sys.argv[1] == "TEST_JIRA":
            metaF.testJira()
        elif sys.argv[1] == "TEST_POLARION_FOR_LEXUS":
            metaF.importExcelToPolarion()
        elif sys.argv[1] == "ADD_INFO_FORD":
            metaF.add_schedule_FORD()
        elif sys.argv[1] == "UPD_INFO_FORD":
            metaF.upd_funcID_FORD()
        elif sys.argv[1] == "DIFF_WORD":
            metaF.diffWordFile()
        elif sys.argv[1] == "UPD_REDMINE":
            metaF.updateiAutoRedmineSubject()
        elif sys.argv[1] == "REDMINE_FOR_24MM":
            metaF.updateTeslinRedmineSubject()
        elif sys.argv[1] == "REDMINE_24MM_HXY":
            metaF.getTeslinRedmineSubject()
        elif sys.argv[1] == "EXTRACT_WORD":
            metaF.ExtractWordFile()
        else:

            metaF.testRedmineApi()
    elif len(sys.argv) == 3:
        if sys.argv[1] == "TEST_PDF":
            metaF.extract_pdf(sys.argv[2])
        if sys.argv[1] == "TEST_RFQ_FORD":
            metaF.extractOutline_rfq_FORD(sys.argv[2])
        if sys.argv[1] == "TEST_RFQ_FORD":
            metaF.copyWordfile_FORD(sys.argv[2])
        if sys.argv[1] == "UPD_TICKET":
            metaF.updateTeslinRedmineTickets(sys.argv[2])

    elif len(sys.argv) == 4:
        if sys.argv[1] == "PDF_TO_IMG":
            pdf_img_path = sys.argv[3] + "\\"+ sys.argv[2].replace("pdf\\","").replace(".pdf","")
            #if not os.path.exists(pdf_img_path):
            #    os.mkdir(pdf_img_path)
            resultSheet =metaF.resultBook.ActiveSheet
            metaF.grabImage_text(sys.argv[2], pdf_img_path,resultSheet,1)
